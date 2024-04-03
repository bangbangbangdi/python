# -------------------- Word操作 --------------------
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.document import Document as Doc


# -------------------- 写操作 --------------------
def create_doc(filename):
    # 创建Doc对象
    document = Document()
    # 添加标题, 数字0是标题的意思 level
    document.add_heading("Kino's Journey", 0)

    # 添加段落
    p = document.add_paragraph("無理矢理他の人と")
    # 声明一个列表
    bold_list = []
    # 往段落中追加内容,并返回追加的对象 我们可以通过修改该对象的属性值来决定文本的样式
    run = p.add_run("同じように")
    # 因为我想要这些文本的样式都统一是 加粗 + 绿色; 如果每次都逐个设置的话就会出现很多重复代码
    # 所以先将他们收集到列表中,最后统一进行修改
    bold_list.append(run)
    p.add_run('大人になるのではなく、')
    # 将需要修改的对象加入列表
    bold_list.append(p.add_run('速度'))
    p.add_run('や')
    bold_list.append(p.add_run('順番'))
    p.add_run('はバラバラでも、')
    bold_list.append(p.add_run('自分自身'))
    p.add_run('納得する方法で、自分自身納得する、納得できる')
    bold_list.append(p.add_run('大人になる'))
    p.add_run('。')

    # 遍历bold_list 对每一个对象都进行 加粗 + 修改颜色的操作
    for run in bold_list:
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 139, 87)

    # 添加一级标题
    document.add_heading('人の痛みが分かる国')
    # 添加带样式的段落
    document.add_paragraph('多数決の国', style='Intense Quote')
    # 添加无序列表
    document.add_paragraph('レールの上の三人の男', style='List Bullet')
    document.add_paragraph('コロシアム', style='List Bullet')
    # 添加有序列表
    document.add_paragraph('大人の国', style='List Number')
    document.add_paragraph('平和な国', style='List Number')
    # 添加图片
    document.add_picture('./testFile/kino1.png', width=Cm(13))

    # 添加分页符
    document.add_section()

    records = (('kino', 'female', '16'), ('erms', 'unknown', '3'), ('shizu', 'male', '20'))
    # 添加表格
    table = document.add_table(rows=1, cols=3)
    table.style = 'Dark List'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'gender'
    hdr_cells[2].text = 'age'
    # 为表格添加新行
    for name, gender, age in records:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = gender
        row_cells[2].text = age

    # 保存文件
    document.save(filename)


# -------------------- 读操作 --------------------
def read_doc(filename):
    doc = Document(filename)  # type: Doc
    for no, p in enumerate(doc.paragraphs):
        print(no, p.text)


# -------------------- 根据模版生成内容 --------------------
def change_doc(template_name):
    employees = [
        {
            'name': 'kino',
            'id': '439943994399',
            'sdate': '2023-11-03',
            'edate': '9999-12-32',
            'department': '环球探险',
            'position': '旅行家',
            'company': '还没想好叫啥名字有限公司'
        }
    ]
    # 对列表进行循环遍历,批量生成Word文档
    for emp_dict in employees:
        doc = Document(template_name)  # type: Doc
        for p in doc.paragraphs:
            if '{' not in p.text:
                continue
            for run in p.runs:
                if '{' not in run.text:
                    continue
                start, end = run.text.find('{'), run.text.find('}')
                key, place_holder = run.text[start + 1:end], run.text[start:end + 1]
                run.text = run.text.replace(place_holder, emp_dict[key])
        doc.add_picture('./testFile/kino2.png', width=Cm(13))
        doc.save(f'./testFile/{emp_dict["name"]}离职证明.docx')


def main():
    # create_doc('./testFile/Kino.docx')
    # read_doc('./testFile/离职证明.docx')
    change_doc('./testFile/离职证明模板.docx')
    pass


if __name__ == '__main__':
    main()
